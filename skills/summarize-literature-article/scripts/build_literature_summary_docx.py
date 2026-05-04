#!/usr/bin/env python
"""Build a Word report from a structured Markdown literature summary."""

from __future__ import annotations

import argparse
import re
from datetime import datetime
from pathlib import Path


def _require_docx():
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except ImportError as exc:
        raise RuntimeError("Word output requires the python-docx package.") from exc
    return Document, WD_ALIGN_PARAGRAPH, Cm, Pt


def _is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def _split_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|") and lines[index].strip().endswith("|"):
        if not _is_table_separator(lines[index]):
            rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
        index += 1
    return rows, index


def _set_normal_style(document, font_size, pt_factory):
    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = font_size
    style.paragraph_format.space_after = pt_factory(6)
    style.paragraph_format.line_spacing = 1.15


def _add_markdown(document, markdown: str):
    lines = markdown.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            rows, index = _split_table(lines, index)
            if rows:
                width = max(len(row) for row in rows)
                table = document.add_table(rows=len(rows), cols=width)
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(width):
                        table.cell(r_idx, c_idx).text = row[c_idx] if c_idx < len(row) else ""
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = min(len(heading.group(1)), 4)
            document.add_heading(heading.group(2).strip(), level=level)
            index += 1
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            document.add_paragraph(bullet.group(1).strip(), style="List Bullet")
            index += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            document.add_paragraph(numbered.group(1).strip(), style="List Number")
            index += 1
            continue

        document.add_paragraph(stripped)
        index += 1


def main() -> int:
    Document, WD_ALIGN_PARAGRAPH, Cm, Pt = _require_docx()

    parser = argparse.ArgumentParser(description="Convert a literature summary Markdown file to Word.")
    parser.add_argument("--markdown", required=True, type=Path, help="Structured Markdown summary")
    parser.add_argument("--out", required=True, type=Path, help="Output .docx path")
    parser.add_argument("--title", default="文献总结报告", help="Document title")
    parser.add_argument("--source", default="", help="Optional source article path or citation")
    args = parser.parse_args()

    markdown_path = args.markdown.resolve()
    if not markdown_path.exists():
        raise FileNotFoundError(markdown_path)

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    _set_normal_style(document, Pt(10.5), Pt)

    title = document.add_heading(args.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(9)
    if args.source:
        src = document.add_paragraph()
        src.alignment = WD_ALIGN_PARAGRAPH.CENTER
        src_run = src.add_run(f"来源：{args.source}")
        src_run.font.size = Pt(9)

    document.add_paragraph("")
    _add_markdown(document, markdown_path.read_text(encoding="utf-8"))

    out_path = args.out.resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.core_properties.title = args.title
    document.core_properties.subject = "Structured literature summary"
    document.save(str(out_path))
    print(str(out_path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
