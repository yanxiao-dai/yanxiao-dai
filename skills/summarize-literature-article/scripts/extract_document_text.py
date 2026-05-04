#!/usr/bin/env python
"""Extract readable text and basic metadata from PDF or DOCX articles."""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any


def _clean_text(text: str | None) -> str:
    if not text:
        return ""
    lines = [line.rstrip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    return "\n".join(lines).strip()


def extract_pdf(path: Path, max_pages: int | None = None) -> tuple[str, dict[str, Any]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF extraction requires the pypdf package.") from exc

    reader = PdfReader(str(path))
    pages = reader.pages[:max_pages] if max_pages else reader.pages
    chunks: list[str] = []
    for index, page in enumerate(pages, start=1):
        text = _clean_text(page.extract_text())
        chunks.append(f"\n\n<!-- Page {index} -->\n\n{text}" if text else f"\n\n<!-- Page {index}: no extractable text -->")

    raw_meta = reader.metadata or {}
    metadata = {
        "file": str(path),
        "format": "pdf",
        "pages_total": len(reader.pages),
        "pages_extracted": len(pages),
        "title": str(raw_meta.get("/Title", "") or ""),
        "author": str(raw_meta.get("/Author", "") or ""),
        "subject": str(raw_meta.get("/Subject", "") or ""),
        "creator": str(raw_meta.get("/Creator", "") or ""),
        "producer": str(raw_meta.get("/Producer", "") or ""),
    }
    return "\n".join(chunks).strip(), metadata


def _docx_table_to_markdown(table: Any) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([_clean_text(cell.text).replace("\n", " ") for cell in row.cells])
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = "| " + " | ".join(normalized[0]) + " |"
    sep = "| " + " | ".join(["---"] * width) + " |"
    body = ["| " + " | ".join(row) + " |" for row in normalized[1:]]
    return "\n".join([header, sep, *body])


def extract_docx(path: Path) -> tuple[str, dict[str, Any]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX extraction requires the python-docx package.") from exc

    doc = Document(str(path))
    chunks: list[str] = []
    for paragraph in doc.paragraphs:
        text = _clean_text(paragraph.text)
        if not text:
            continue
        style = (paragraph.style.name or "").lower() if paragraph.style else ""
        if style.startswith("heading"):
            level = "".join(ch for ch in style if ch.isdigit()) or "2"
            level_num = min(max(int(level), 1), 6)
            chunks.append(f"{'#' * level_num} {text}")
        else:
            chunks.append(text)

    for idx, table in enumerate(doc.tables, start=1):
        table_md = _docx_table_to_markdown(table)
        if table_md:
            chunks.append(f"\n\nTable {idx}\n\n{table_md}")

    props = doc.core_properties
    metadata = {
        "file": str(path),
        "format": "docx",
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "title": props.title or "",
        "author": props.author or "",
        "subject": props.subject or "",
        "keywords": props.keywords or "",
        "created": props.created.isoformat() if props.created else "",
        "modified": props.modified.isoformat() if props.modified else "",
    }
    return "\n\n".join(chunks).strip(), metadata


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract text from a journal article PDF or DOCX.")
    parser.add_argument("input", type=Path, help="Path to .pdf or .docx article")
    parser.add_argument("--out", type=Path, help="Markdown output path")
    parser.add_argument("--metadata", type=Path, help="JSON metadata output path")
    parser.add_argument("--max-pages", type=int, default=None, help="Optional maximum pages to extract from PDF")
    args = parser.parse_args()

    input_path = args.input.resolve()
    if not input_path.exists():
        raise FileNotFoundError(input_path)

    suffix = input_path.suffix.lower()
    if suffix == ".pdf":
        text, metadata = extract_pdf(input_path, args.max_pages)
    elif suffix == ".docx":
        text, metadata = extract_docx(input_path)
    elif suffix == ".doc":
        raise RuntimeError("Legacy .doc files must be converted to .docx before extraction.")
    else:
        raise RuntimeError(f"Unsupported input format: {suffix}")

    metadata["characters"] = len(text)
    metadata["likely_scanned_or_empty"] = len(text.strip()) < 500

    out_path = args.out or input_path.with_suffix(".extracted.md")
    meta_path = args.metadata or input_path.with_suffix(".metadata.json")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    meta_path.parent.mkdir(parents=True, exist_ok=True)

    out_path.write_text(text + "\n", encoding="utf-8")
    meta_path.write_text(json.dumps(metadata, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    print(json.dumps({"text": str(out_path), "metadata": str(meta_path), **metadata}, ensure_ascii=False))
    if metadata["likely_scanned_or_empty"]:
        print("Warning: extracted text is short; OCR may be required.", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
